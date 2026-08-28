"""Structural and render QA for the real multi-format RAG corpus."""

from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageChops


ROOT = Path(__file__).resolve().parents[1]
CORPUS = ROOT / "output" / "multitype-corpus"


def zip_check(path: Path) -> dict:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        names = archive.namelist()
    return {"valid_zip": bad is None, "member_count": len(names), "bad_member": bad}


def docx_check(path: Path) -> dict:
    result = zip_check(path)
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    section = document.sections[0]
    table = document.tables[0]
    result.update(
        {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "rows_in_first_table": len(table.rows),
            "page_inches": [round(section.page_width.inches, 2), round(section.page_height.inches, 2)],
            "margins_inches": [
                round(section.top_margin.inches, 2),
                round(section.right_margin.inches, 2),
                round(section.bottom_margin.inches, 2),
                round(section.left_margin.inches, 2),
            ],
            "contains_fact": "17 分钟" in text,
            "contains_question": "异常外联隔离必须在多少分钟内完成" in text,
        }
    )
    return result


def pptx_check(path: Path) -> dict:
    result = zip_check(path)
    with zipfile.ZipFile(path) as archive:
        slide_names = sorted(
            name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
        )
        xml_text = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in slide_names)
    renders = []
    for image_path in sorted((CORPUS / "rendered-pptx").glob("slide-*.png")):
        image = Image.open(image_path).convert("RGB")
        background = Image.new("RGB", image.size, image.getpixel((0, 0)))
        bbox = ImageChops.difference(image, background).getbbox()
        touches_edge = False
        if bbox:
            touches_edge = bbox[0] <= 1 or bbox[1] <= 1 or bbox[2] >= image.width - 1 or bbox[3] >= image.height - 1
        renders.append(
            {
                "file": image_path.name,
                "size": list(image.size),
                "content_bbox": list(bbox) if bbox else None,
                "content_touches_edge": touches_edge,
            }
        )
    result.update(
        {
            "slides": len(slide_names),
            "contains_fact": "23 minutes" in xml_text,
            "contains_question": "How quickly must PPTX-7319" in xml_text,
            "rendered_slides": renders,
        }
    )
    return result


def xlsx_check(path: Path) -> dict:
    result = zip_check(path)
    with zipfile.ZipFile(path) as archive:
        worksheet_names = [name for name in archive.namelist() if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name)]
        payload = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in archive.namelist() if name.endswith(".xml"))
    result.update(
        {
            "worksheets": len(worksheet_names),
            "contains_high_value": "HIGH" in payload,
            "contains_report_id_2": ">2<" in payload,
        }
    )
    return result


def pdf_check(path: Path) -> dict:
    data = path.read_bytes()
    source_png = ROOT / "output" / "smoke" / "mindbridge_mineru_ocr_utf8.png"
    image = Image.open(source_png)
    return {
        "valid_header": data.startswith(b"%PDF-"),
        "has_eof": b"%%EOF" in data[-2048:],
        "approx_pages": len(re.findall(rb"/Type\s*/Page(?!s)", data)),
        "source_image_size": list(image.size),
    }


def main() -> int:
    report = {
        "pdf": pdf_check(CORPUS / "ocr-incident-report.pdf"),
        "docx": docx_check(CORPUS / "incident-isolation-brief.docx"),
        "pptx": pptx_check(CORPUS / "recovery-drill-brief.pptx"),
        "xlsx": xlsx_check(CORPUS / "mindbridge-risk-ledger.xlsx"),
        "markdown": {
            "contains_fact": "35%" in (CORPUS / "release-notes.md").read_text(encoding="utf-8"),
        },
        "text": {
            "contains_fact": "42 秒" in (CORPUS / "oncall-note.txt").read_text(encoding="utf-8"),
        },
    }
    output = CORPUS / "validation-report.json"
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    failures = []
    for file_type, checks in report.items():
        if isinstance(checks, dict):
            for name, value in checks.items():
                if name.startswith(("valid_", "contains_", "has_")) and value is not True:
                    failures.append(f"{file_type}.{name}")
    if any(item["content_touches_edge"] for item in report["pptx"]["rendered_slides"]):
        failures.append("pptx.render_content_touches_edge")
    if len(report["pptx"]["rendered_slides"]) != report["pptx"]["slides"]:
        failures.append("pptx.render_count")
    if failures:
        print("FAILURES:", failures)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
