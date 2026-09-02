"""P2：多格式渲染 Drive（生成真实文件并做结构校验）。

用法:
    python -m scripts.enterprise_rag.render --scale S1 --seed 20260828
"""
from __future__ import annotations

import json
from pathlib import Path

from scripts.enterprise_rag.config import RunConfig
from scripts.enterprise_rag.renderers import content, render_all


def validate_binary_structure(root: Path) -> list[dict]:
    """对真实二进制文件做打开/结构校验（计划 §8.2 / §8 视觉检查记录）。"""
    checks: list[dict] = []
    for path in root.rglob("*"):
        if path.suffix in (".docx", ".xlsx", ".pptx", ".pdf"):
            ok = False
            try:
                if path.suffix == ".docx":
                    from docx import Document
                    d = Document(str(path))
                    ok = len(d.paragraphs) >= 2 or len(d.tables) >= 1
                elif path.suffix == ".xlsx":
                    from openpyxl import load_workbook
                    wb = load_workbook(str(path), read_only=True)
                    ws = wb.active
                    ok = ws.max_row and ws.max_row >= 1
                elif path.suffix == ".pptx":
                    from pptx import Presentation
                    prs = Presentation(str(path))
                    ok = len(prs.slides) >= 1
                elif path.suffix == ".pdf":
                    from pypdf import PdfReader
                    r = PdfReader(str(path))
                    ok = len(r.pages) >= 1
            except Exception as exc:  # noqa: BLE001
                checks.append({"path": str(path), "ok": False, "error": str(exc)[:200]})
                continue
            checks.append({"path": str(path), "ok": ok})
    return checks


def render_scale(scale: str, seed: int) -> dict:
    cfg = RunConfig(run_id=f"render-{scale}-{seed}", scale=scale, seed=seed)
    docs, faq_by_product, catalog, facts = content.build_docs(scale, seed)
    files = cfg.files_dir
    # 清空生成目录，避免陈旧文件污染（计划 §8：每次渲染以当前 plan 为准）
    if files.exists():
        for p in sorted(files.rglob("*"), reverse=True):
            p.unlink() if p.is_file() else p.rmdir()
    files.mkdir(parents=True, exist_ok=True)
    manifest = render_all(docs, faq_by_product, files)
    bin_checks = validate_binary_structure(files)
    manifest["binary_checked"] = len(bin_checks)
    manifest["binary_ok"] = sum(1 for c in bin_checks if c["ok"])
    manifest["binary_failed"] = [c for c in bin_checks if not c["ok"]][:40]
    (cfg.out_dir).mkdir(parents=True, exist_ok=True)
    (cfg.out_dir / "render-report.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--scale", default="S1")
    ap.add_argument("--seed", type=int, default=20260828)
    a = ap.parse_args()
    m = render_scale(a.scale, a.seed)
    print("documents:", m["documents"], "faq:", m["faq_entries"], "bytes:", m["total_bytes"])
    print("formats:", m["by_format"])
    print("binary ok:", m.get("binary_ok"), "/", m.get("binary_checked"),
          "failed:", len(m.get("binary_failed", [])))