"""Create deterministic DOCX, Markdown and TXT fixtures for live RAG validation."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "multitype-corpus"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def cell_geometry(cell, width, fill=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW")) or OxmlElement("w:tcW")
    if tc_w.getparent() is None:
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    margins = tc_pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if margins.getparent() is None:
        tc_pr.append(margins)
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = margins.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}")
        if node.getparent() is None:
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    if fill:
        shade = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
        if shade.getparent() is None:
            tc_pr.append(shade)
        shade.set(qn("w:fill"), fill)


def create_docx(path):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("RAG 多类型语料验证｜内部"), 9, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("MindBridge 文档处理验证样本"), 9, color="6B7280")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("应急隔离操作简报"), 23, True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("文档编号 DOCX-4821｜真实 DOCX 解析与切块验证"), 12, color="374151")

    for label, value in (("适用范围", "生产环境异常连接隔离"), ("责任角色", "安全值班经理"), ("密级", "INTERNAL")):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(label + "："), bold=True)
        set_font(paragraph.add_run(value))

    doc.add_heading("关键要求", level=1)
    paragraph = doc.add_paragraph()
    set_font(paragraph.add_run("检测到异常外联后，隔离动作必须在 17 分钟内完成；"), bold=True, color="9B1C1C")
    set_font(paragraph.add_run("随后由安全值班经理在工单中登记证据链。"))

    doc.add_heading("执行步骤", level=1)
    for number, step in enumerate((
        "确认告警指纹与资产归属，避免误隔离共享出口。",
        "执行网络隔离并保存策略变更编号、时间戳和操作者。",
        "验证异常连接归零，将结果写入 DOCX-4821 处置记录。",
    ), 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        set_font(paragraph.add_run(f"{number}. "), bold=True, color="2E74B5")
        set_font(paragraph.add_run(step))

    doc.add_heading("升级矩阵", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment, table.autofit, table.style = WD_TABLE_ALIGNMENT.LEFT, False, "Table Grid"
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    widths = (2160, 3600, 3600)
    for index, value in enumerate(("级别", "触发条件", "升级对象")):
        cell = table.rows[0].cells[index]
        cell_geometry(cell, widths[index], "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cell.paragraphs[0].add_run(value), bold=True)
    for row in (("P1", "隔离超过 17 分钟", "安全负责人"), ("P2", "证据链字段不完整", "值班经理")):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell_geometry(cells[index], widths[index])
            set_font(cells[index].paragraphs[0].add_run(value))

    doc.add_heading("验收口径", level=1)
    paragraph = doc.add_paragraph()
    set_font(paragraph.add_run("检索问题："), bold=True)
    set_font(paragraph.add_run("异常外联隔离必须在多少分钟内完成？"))
    paragraph = doc.add_paragraph()
    set_font(paragraph.add_run("预期答案："), bold=True)
    set_font(paragraph.add_run("17 分钟。"))
    doc.save(path)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    create_docx(OUTPUT / "incident-isolation-brief.docx")
    (OUTPUT / "release-notes.md").write_text(
        "# 检索网关发布说明\n\n## 灰度规则\n\n"
        "版本 `MD-2608` 的灰度流量上限为 **35%**。当错误率连续两个窗口高于 1.5% 时，必须回滚到上一代索引。\n\n"
        "## 查询验证\n\n- 问题：MD-2608 的灰度流量上限是多少？\n- 预期答案：35%。\n",
        encoding="utf-8",
    )
    (OUTPUT / "oncall-note.txt").write_text(
        "值班交接编号 TXT-9064。\n缓存雪崩告警触发后，首轮健康检查间隔固定为 42 秒。\n"
        "若连续三次失败，则切换只读降级并通知平台负责人。\n检索问题：首轮健康检查间隔是多少秒？预期答案：42 秒。\n",
        encoding="utf-8",
    )
    print(OUTPUT)


if __name__ == "__main__":
    main()
